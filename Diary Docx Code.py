import os
from docx import Document  # Import the Document class from python-docx

Date = input('Enter Date:')
Diary_Entry = input('Tell me all about your day bestie: ')

file_path = r"C:\Users\KING G\OneDrive\Desktop\git\Python\diary.docx"


try:
    # Create a new Document object
    doc = Document()

    # Add the Date and Diary_Entry to the document
    doc.add_paragraph(f'Date: {Date}')
    doc.add_paragraph(Diary_Entry)

    # Save the document to the specified file path
    doc.save(file_path)

    print(f'Diary entry filled and saved to {file_path} successfully.')
except Exception as e:
    print(f'An error was encountered: {e}')
